from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOWNLOADS = Path(r"C:\Users\bibe\Downloads")
UNIFIED_SRC = DOWNLOADS / "AquaAI_Unified_Technical_Documentation_v2.docx"
USER_GUIDE_SRC = DOWNLOADS / "AquaAI_Intelligence_Complete_User_Guide (1).docx"
UNIFIED_OUT = DOWNLOADS / "AquaAI_Unified_Technical_Documentation_v2_1_Phase5.docx"
USER_GUIDE_OUT = DOWNLOADS / "AquaAI_Intelligence_Complete_User_Guide_v1_1_Phase5.docx"
BACKEND_GUIDE_OUT = DOWNLOADS / "AquaAI_Backend_Developer_Branch_Handoff_Guide.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def get_existing_style(doc: Document, style_name: str):
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.name == style_name:
            return paragraph.style
    return None


def get_body_style(doc: Document):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text and not style_name.startswith("Heading") and style_name != "Title":
            return paragraph.style
    return None


def insert_block_before(paragraph, items: list[tuple[object | None, str]]) -> list:
    inserted = []
    for style, text in items:
        new_p = paragraph.insert_paragraph_before(text)
        if style:
            new_p.style = style
        inserted.append(new_p)
    return inserted


def format_body_paragraph(paragraph, font_size: int = 11) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        run.bold = False
        run.font.size = Pt(font_size)


def ensure_spacing(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        fmt = paragraph.paragraph_format
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(4)
        else:
            fmt.space_after = Pt(4)


def deduplicate_named_styles(docx_path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    val_attr = "{%s}val" % ns["w"]
    style_id_attr = "{%s}styleId" % ns["w"]

    with ZipFile(docx_path, "r") as src_zip:
        styles_xml = src_zip.read("word/styles.xml")
        root = ET.fromstring(styles_xml)
        seen: set[tuple[str, str]] = set()
        for style in list(root.findall("w:style", ns)):
            style_id = style.get(style_id_attr, "")
            name = style.find("w:name", ns)
            name_val = name.get(val_attr, "") if name is not None else ""
            key = (style_id, name_val)
            if key in seen:
                root.remove(style)
            else:
                seen.add(key)

        updated_styles_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_path = Path(tmp_file.name)

        with ZipFile(tmp_path, "w", compression=ZIP_DEFLATED) as dst_zip:
            for info in src_zip.infolist():
                data = updated_styles_xml if info.filename == "word/styles.xml" else src_zip.read(info.filename)
                dst_zip.writestr(info, data)

    tmp_path.replace(docx_path)


def update_unified_doc() -> None:
    doc = Document(UNIFIED_SRC)
    heading2 = get_existing_style(doc, "Heading 2")
    body_style = get_body_style(doc)
    intelligence_context_paragraph = None

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if "Unified Technical Documentation v2.0" in paragraph.text:
                paragraph.text = paragraph.text.replace("v2.0", "v2.1")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.bold = False
                    run.font.size = Pt(8)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Version 2.0 — Consolidated & Locked":
            paragraph.text = "Version 2.1 — Consolidated, Locked & Phase 5 Integration Updated"
        elif text == "Date: April 2026":
            paragraph.text = "Date: 30 April 2026"
        elif text.startswith("This document consolidates the original AquaAI Technical Documentation"):
            paragraph.text = (
                "This document consolidates the original AquaAI Technical Documentation "
                "(March 2025 v1.0) with the updated subscription, promotional pricing, "
                "referral architecture, intelligence layer additions, and the Phase 5 "
                "backend-to-intelligence integration changes completed on 30 April 2026. "
                "It supersedes all prior technical drafts."
            )
        elif text == "input → Django API → Aqua Intelligence Brain → LLM → JSON → Frontend":
            paragraph.text = "input → Django API → Aqua Intelligence Brain → LLM → JSON → Frontend"
        elif text.startswith("The Intelligence Brain enriches every LLM prompt"):
            intelligence_context_paragraph = paragraph
            paragraph.text = (
                "The Intelligence Brain enriches every LLM prompt with contextual data drawn "
                "from the ai_* intelligence tables: trust trajectory, churn risk, off-platform "
                "flags, sentiment, and entity baselines. The LLM never reasons in isolation. "
                "Phase 5 also formalised the reverse event path: badge, maintenance, and "
                "messaging events are dispatched post-commit into the Intelligence API so "
                "learning signals are based only on committed operational data."
            )
        elif text.startswith("badges_trustscoresnapshot: Historical snapshots of user trust levels over time"):
            paragraph.text = (
                "badges_trustscoresnapshot: Historical snapshots of user trust levels over time "
                "— immutable evidence chain. Phase 5 stores local_baseline, "
                "intelligence_modifier, final_trust_score, and intelligence context inside "
                "contributing_factors JSON rather than introducing new database columns."
            )
        elif text.startswith("ai_recommendation_outcomes: Recommendation tracking and feedback loop."):
            paragraph.text = (
                "ai_recommendation_outcomes: Recommendation tracking and feedback loop. "
                "Phase 5 aligns writes to the production schema using source, status, "
                "created_at, resolved_at, action_taken, was_helpful, and outcome_score."
            )
        elif text.startswith("ai_learning_signals: Continuous learning inputs for model refinement."):
            paragraph.text = (
                "ai_learning_signals: Continuous learning inputs for model refinement. "
                "Phase 5 standardises recommendation, badge, sentiment, and booking signals "
                "to include aquaai_user_id, entity_id, entity_type, and recommendation linkage "
                "where available."
            )
        elif text == "10.1 Current Version (v2.0 — April 2026)":
            paragraph.text = "10.1 Current Version (v2.1 — 30 April 2026)"
        elif text == "Version: 2.0":
            paragraph.text = "Version: 2.1"
        elif text == "Date of Issue: April 2026.":
            paragraph.text = "Date of Issue: 30 April 2026."

    if intelligence_context_paragraph is not None:
        inserted = insert_block_before(
            intelligence_context_paragraph,
            [
                (
                    body_style,
                    "Phase 5 adds a reverse event path: committed platform events now flow "
                    "back from Django to Intelligence through the authenticated "
                    "/intelligence/events/ingest bridge.",
                ),
            ],
        )
        for paragraph in inserted:
            format_body_paragraph(paragraph)

    doc_control = next(p for p in doc.paragraphs if p.text.strip() == "Document Control")
    inserted = insert_block_before(
        doc_control,
        [
            (heading2, "10.5 Phase 5 Integration Update (30 April 2026)"),
            (
                body_style,
                "- Django remains the authoritative operational layer for trust, while "
                "Intelligence now contributes a bounded modifier and contextual reasoning "
                "signal rather than replacing the backend score.",
            ),
            (
                body_style,
                "- Trust snapshots continue to use the existing badges_trustscoresnapshot "
                "schema. The local baseline, intelligence modifier, final score, and "
                "intelligence context are stored inside contributing_factors JSON, so "
                "this release introduces zero new database migrations.",
            ),
            (
                body_style,
                "- The backend now pushes committed badge, maintenance, and messaging events "
                "to POST /intelligence/events/ingest using transaction.on_commit(), ensuring "
                "Intelligence only learns from committed operational rows.",
            ),
            (
                body_style,
                "- The Intelligence ingest endpoint is authenticated with the shared service "
                "JWT. Unauthenticated callers are rejected, while authenticated backend calls "
                "can trigger off-platform detection, sentiment analysis, recommendation "
                "outcome tracking, and learning-signal emission.",
            ),
            (
                body_style,
                "- AI care tasks and tank recommendations now carry recommendation IDs from the "
                "Intelligence layer through backend persistence and back into the recommendation "
                "outcome loop when a task is completed or missed.",
            ),
            (
                body_style,
                "- Schema alignment fixes now match the live tables for ai_recommendation_outcomes, "
                "ai_learning_signals, ai_off_platform_detections, ai_trust_score_history, "
                "badges_userbadge, badges_badgedefinition, marketplace_marketplacelisting, and "
                "breeders_breederstock.",
            ),
            (
                body_style,
                "- Gamification and micro-segmentation no longer depend on brittle compatibility "
                "view columns such as reviewee_id or species_name. They now fall back to "
                "schema-aligned base tables and authoritative review/listing sources.",
            ),
            (
                body_style,
                "- Operational sequencing is locked for deployment: Intelligence branch first, "
                "backend branch second, with production verification gates around authenticated "
                "event ingest, live trust modifier reads, and post-commit smoke tests.",
            ),
        ],
    )
    for paragraph in inserted[1:]:
        format_body_paragraph(paragraph)

    ensure_spacing(doc)
    doc.save(UNIFIED_OUT)
    deduplicate_named_styles(UNIFIED_OUT)


def update_user_guide() -> None:
    doc = Document(USER_GUIDE_SRC)
    heading2 = get_existing_style(doc, "Heading 2")
    body_style = get_body_style(doc)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Version 1.0  |  April 2026":
            paragraph.text = "Version 1.1  |  30 April 2026"
        elif text.startswith("The mobile app and provider dashboard never communicate directly"):
            paragraph.text = (
                "The mobile app and provider dashboard never communicate directly with the "
                "Intelligence API. All requests flow through the Django backend, which "
                "authenticates the user, validates ownership, and enriches the response with "
                "intelligence outputs. Phase 5 also introduced an authenticated event bridge "
                "back into Intelligence so committed messages, badge changes, and maintenance "
                "outcomes can improve the learning loop without exposing the Intelligence API "
                "to end users."
            )
        elif text.startswith("Care tasks are dynamic"):
            paragraph.text = (
                "Care tasks are dynamic — they adapt based on your tank’s current state. If a "
                "water anomaly is detected, emergency tasks are generated immediately (such as "
                "“Perform 50% water change — ammonia spike detected”). Routine tasks also shift "
                "based on seasonality, stocking changes, and your historical completion patterns."
            )
        elif text.startswith("Your care task completion rate directly influences your trust score."):
            paragraph.text = (
                "Your care task completion rate directly influences your trust score. "
                "Consistently completing tasks on time demonstrates responsible fishkeeping and "
                "improves your standing in the AquaAI ecosystem. Phase 5 now tracks the "
                "outcome of AI-generated recommendations end-to-end, so completed or missed "
                "tasks teach the platform which recommendations are most useful, most urgent, "
                "and most likely to help similar tanks."
            )
        elif text.startswith("Your Trust Score is the cornerstone of your professional reputation"):
            paragraph.text = (
                "Your Trust Score is the cornerstone of your professional reputation on AquaAI. "
                "It is a dynamic, AI-calculated metric ranging from 0 to 100 that reflects your "
                "reliability, expertise, responsiveness, and client satisfaction. The score shown "
                "to users now comes from AquaAI’s authoritative backend trust snapshot and is "
                "then enriched with intelligence context such as trajectory, churn risk, and "
                "off-platform indicators."
            )
        elif text.startswith("The trust score is calculated from multiple signals including your booking completion rate"):
            paragraph.text = (
                "The trust score is calculated from multiple signals including your booking "
                "completion rate, average response time to inquiries, client review sentiment, "
                "platform compliance, care task quality (if applicable), and consistent platform "
                "engagement. The backend remains the source of truth for the final score, while "
                "the Intelligence layer contributes a controlled modifier and explanatory context."
            )
        elif text.startswith("Each badge contributes points to your trust score, creating a virtuous cycle"):
            paragraph.text = (
                "Each badge contributes points to your trust score, creating a virtuous cycle: "
                "good performance earns badges, which boost your score, which increases your "
                "visibility, which brings more clients. Consultant and breeder badge progress "
                "now uses corrected review and listing sources, so progress nudges and milestone "
                "estimates reflect live platform activity more accurately."
            )
        elif text.startswith("The system distinguishes between different types of off-platform signals"):
            paragraph.text = (
                "The system distinguishes between different types of off-platform signals, "
                "including payment redirect attempts, external contact sharing, platform bypass "
                "language, and discount-for-bypass offers. Each type carries a different risk "
                "weight. Phase 5 tightened the processing path so analysis only runs after a "
                "message is successfully committed, avoiding false learning from rolled-back or "
                "failed sends."
            )

    chapter2 = next(p for p in doc.paragraphs if p.text.strip() == "Chapter 2: Tank Owner Guide")
    inserted = insert_block_before(
        chapter2,
        [
            (heading2, "April 2026 Platform Update"),
            (
                body_style,
                "The April 2026 Phase 5 update tightened the collaboration between the AquaAI "
                "backend and the Intelligence service. The result is more accurate trust scores, "
                "more reliable task feedback, and safer communication monitoring across the app.",
            ),
            (
                body_style,
                "- Trust scores now display the authoritative backend snapshot, with intelligence "
                "signals layered in as contextual modifiers rather than as a separate competing score.",
            ),
            (
                body_style,
                "- Every in-platform message can now trigger off-platform safety analysis and "
                "sentiment scoring after the message is successfully sent.",
            ),
            (
                body_style,
                "- AI care tasks and tank recommendations are now tracked end-to-end, so task "
                "completion and misses feed back into future recommendation quality.",
            ),
            (
                body_style,
                "- Consultant and breeder badge progress, cohort benchmarking, and behavioral "
                "profiling now use corrected review and listing data sources for better accuracy.",
            ),
        ],
    )
    for paragraph in inserted[1:]:
        format_body_paragraph(paragraph)

    glossary_end = next(p for p in doc.paragraphs if p.text.strip() == "— End of Guide —")
    inserted = insert_block_before(
        glossary_end,
        [
            (
                body_style,
                "Recommendation Learning Loop — The feedback system that tracks whether an "
                "AI-generated care task or recommendation was completed, missed, helpful, or "
                "dismissed so future recommendations can improve over time.",
            ),
            (
                body_style,
                "Recommendation Tracking — The background linkage that carries an AI-generated "
                "task or recommendation from creation, through task scheduling, to the final "
                "outcome recorded by the platform.",
            ),
        ],
    )
    for paragraph in inserted:
        format_body_paragraph(paragraph)

    ensure_spacing(doc)
    doc.save(USER_GUIDE_OUT)
    deduplicate_named_styles(USER_GUIDE_OUT)


def build_backend_handoff_guide() -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AquaAI Backend Developer Branch Handoff Guide")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Phase 5 backend + intelligence integration follow-up\nVersion 1.0 | 30 April 2026"
    )
    subtitle_run.italic = True

    intro = doc.add_paragraph()
    intro.add_run("Purpose: ").bold = True
    intro.add_run(
        "This guide gives the backend developer the exact branch-based workflow for reviewing, "
        "merging, deploying, and verifying the Phase 5 integration work that has already been "
        "pushed to GitHub branches."
    )

    doc.add_paragraph("1. Branch Map", style="Heading 1")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Repository", "Branch", "Base", "Notes"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True)

    rows = [
        (
            "backend_aqua_ai-1",
            "codex/phase5-backend",
            "s3_optimized",
            "Merge after intelligence branch is validated.",
        ),
        (
            "AQUAAI-Inteligence",
            "codex/phase5-intelligence",
            "main",
            "Must be reviewed and deployed first.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    doc.add_paragraph("2. What Changed", style="Heading 1")
    changes = [
        "Hybrid trust scoring now uses Django as the authoritative source while Intelligence contributes a bounded modifier and context.",
        "Committed events now flow from the backend into the authenticated Intelligence ingest endpoint for badge, maintenance, and messaging flows.",
        "AI-generated recommendations now retain recommendation IDs across care-task generation, backend persistence, and outcome resolution.",
        "Schema alignment fixes now match the live production tables for learning signals, recommendation outcomes, off-platform detections, trust history, gamification, and micro-segmentation.",
        "No new database migrations are required for this Phase 5 update.",
    ]
    for item in changes:
        doc.add_paragraph(item, style="List Paragraph")

    doc.add_paragraph("3. Merge and Deploy Order", style="Heading 1")
    order = [
        "Review and merge the intelligence branch first.",
        "Deploy the intelligence service and confirm /intelligence/events/ingest is live and authenticated.",
        "Review and merge the backend branch second.",
        "Deploy the backend service only after the intelligence endpoint and trust modifier endpoint are confirmed healthy.",
    ]
    for index, item in enumerate(order, 1):
        doc.add_paragraph(f"{index}. {item}")

    doc.add_paragraph("4. Required Validation Before Merge", style="Heading 1")
    checks = [
        "Backend branch: run the focused Django test set for trust modifier integration, message intelligence bridging, care-task metadata propagation, and legacy_client logging.",
        "Intelligence branch: run the focused test set for recommendation tracking, event ingest, off-platform schema alignment, authenticated event ingest, and Stream 5B schema alignment.",
        "Confirm that neither repository includes .env, test.sqlite3, or transient local files in the PR.",
        "Confirm that the backend diff does not reintroduce the deleted no-op trust snapshot migration.",
    ]
    for item in checks:
        doc.add_paragraph(item, style="List Paragraph")

    doc.add_paragraph("5. Production Smoke Gates", style="Heading 1")
    smoke_table = doc.add_table(rows=1, cols=3)
    smoke_table.style = "Table Grid"
    for cell, text in zip(smoke_table.rows[0].cells, ["Gate", "What to Verify", "Pass Condition"]):
        set_cell_text(cell, text, bold=True)

    smoke_rows = [
        (
            "Intelligence auth",
            "POST /intelligence/events/ingest without auth and with the service JWT",
            "401 or 422 without auth; 200 with the shared backend JWT",
        ),
        (
            "Trust modifier",
            "Call the live /intelligence/trust/<entity>/<id>/modifier endpoint with the production JWT",
            "200 response with a non-null JSON body",
        ),
        (
            "Backend trust",
            "Recompute a real trust score from Django shell",
            "Returned payload includes local_baseline, intelligence_modifier, and final trust_score",
        ),
        (
            "Message safety",
            "Send a real in-platform smoke-test message",
            "Message commit triggers off-platform detection and sentiment analysis",
        ),
        (
            "Recommendation loop",
            "Complete or miss a tracked AI care task",
            "Recommendation outcome is resolved and linked learning signals are emitted",
        ),
    ]
    for row in smoke_rows:
        cells = smoke_table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    doc.add_paragraph("6. Backend Developer Checklist", style="Heading 1")
    checklist = [
        "Review the intelligence PR first and do not merge the backend PR ahead of it.",
        "Pull the target base branch before reviewing to avoid hidden drift.",
        "Run the exact focused tests documented in the PR descriptions.",
        "Check GitHub branch protection, required reviews, and CI status before merge.",
        "After deploy, run the smoke gates in sequence and capture the results in the release notes or deployment journal.",
        "If any gate fails, stop, rollback the affected service, and investigate before continuing.",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Paragraph")

    doc.add_paragraph("7. Rollback Notes", style="Heading 1")
    rollback = [
        "If the intelligence deploy fails, rollback intelligence first. The backend degrades gracefully when the ingest endpoint is unavailable, but features depending on post-commit learning will pause.",
        "If the backend deploy fails, rollback the backend branch or merge commit. Existing intelligence-side schema-aligned changes remain valid.",
        "Because no new database migrations are introduced in this release, rollback is code-only.",
    ]
    for item in rollback:
        doc.add_paragraph(item, style="List Paragraph")

    ensure_spacing(doc)
    doc.save(BACKEND_GUIDE_OUT)
    deduplicate_named_styles(BACKEND_GUIDE_OUT)


def main() -> None:
    update_unified_doc()
    update_user_guide()
    build_backend_handoff_guide()
    print(UNIFIED_OUT)
    print(USER_GUIDE_OUT)
    print(BACKEND_GUIDE_OUT)


if __name__ == "__main__":
    main()
